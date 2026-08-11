from io import BytesIO

from docx import Document

from app.ai.services.models import GenResumeSchema, ResumeSchema


def _join_non_empty(parts: list[str], separator: str = " | ") -> str:
    return separator.join(part for part in parts if part)


def _render_resume_docx(resume: ResumeSchema | GenResumeSchema) -> bytes:
    document = Document()

    document.add_heading(resume.name, level=0)

    contact = resume.contact
    contact_line = _join_non_empty([
        contact.email,
        contact.phone,
        contact.location,
        contact.linkedin,
        contact.github,
        contact.portfolio,
    ])
    if contact_line:
        document.add_paragraph(contact_line)

    if resume.education:
        document.add_heading("Education", level=1)
        for education in resume.education:
            title = _join_non_empty([
                education.degree,
                f"in {education.major}" if education.major else "",
            ], separator=" ")
            heading = _join_non_empty([title, education.institution])
            if heading:
                paragraph = document.add_paragraph()
                paragraph.add_run(heading).bold = True

            details = _join_non_empty([
                _join_non_empty([education.start_date, education.end_date], separator=" - "),
                f"GPA: {education.gpa}" if education.gpa else "",
            ])
            if details:
                document.add_paragraph(details)

    if resume.experience:
        document.add_heading("Experience", level=1)
        for experience in resume.experience:
            heading = document.add_paragraph()
            heading.add_run(
                _join_non_empty([experience.role, experience.company])
            ).bold = True

            dates = _join_non_empty([experience.start_date, experience.end_date], separator=" - ")
            if dates:
                heading.add_run(f" ({dates})")

            for highlight in experience.highlights:
                if highlight:
                    document.add_paragraph(highlight, style="List Bullet")

    if resume.projects:
        document.add_heading("Projects", level=1)
        for project in resume.projects:
            paragraph = document.add_paragraph()
            paragraph.add_run(project.name).bold = True

            if project.url:
                paragraph.add_run(f" | {project.url}")

            if project.description:
                document.add_paragraph(project.description)

            for highlight in project.highlights:
                if highlight:
                    document.add_paragraph(highlight, style="List Bullet")

    if resume.skills:
        document.add_heading("Skills", level=1)
        for category in resume.skills:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{category.category}: ").bold = True
            paragraph.add_run(", ".join(skill for skill in category.skills if skill))

    if resume.certifications:
        document.add_heading("Certifications", level=1)
        for certification in resume.certifications:
            paragraph = document.add_paragraph()
            paragraph.add_run(certification.name).bold = True

            details = _join_non_empty([
                certification.issuer,
                certification.issue_date,
                certification.url,
            ])
            if details:
                paragraph.add_run(f" | {details}")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_resume_docx(resume: ResumeSchema) -> bytes:
    """Render a parsed ResumeSchema into a DOCX document."""
    return _render_resume_docx(resume)


def build_gen_resume_docx(resume: GenResumeSchema) -> bytes:
    """Render a generated GenResumeSchema into a downloadable DOCX document."""
    return _render_resume_docx(resume)
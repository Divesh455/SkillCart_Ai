import io
import os
import fitz  # PyMuPDF
import docx
from app.core.exceptions import DocumentParsingException

def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text_parts.append(page_text)
        doc.close()
        text = "\n".join(text_parts).strip()
        if not text:
            raise DocumentParsingException("PDF is empty or has no readable text.")
        return text
    except Exception as e:
        if isinstance(e, DocumentParsingException):
            raise e
        raise DocumentParsingException(f"Failed to parse PDF: {str(e)}")

def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        text = "\n".join(full_text).strip()
        if not text:
            raise DocumentParsingException("DOCX file is empty or has no readable text.")
        return text
    except Exception as e:
        if isinstance(e, DocumentParsingException):
            raise e
        raise DocumentParsingException(f"Failed to parse DOCX: {str(e)}")

def extract_text_from_file(filename: str, file_bytes: bytes) -> str:
    """Determine the file type by filename extension and extract its text content."""
    _, ext = os.path.splitext(filename.lower())
    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext in [".docx", ".doc"]:
        return parse_docx(file_bytes)
    else:
        raise DocumentParsingException(f"Unsupported file format: {ext}. Only PDF and DOCX are supported.")
